#!/usr/bin/env python3
"""Generate TECHNICAL_DOCUMENT.docx from the markdown content."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re, os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x07, 0x5F, 0x50)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(12)

# ── Margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BRAND = RGBColor(0x07, 0x5F, 0x50)
GREY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(sh)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '075F50')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True
                r.font.size = Pt(9)
    # Data rows
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

def p(text, bold=False, italic=False, size=None, color=None, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: para.alignment = align
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.5 + level * 1)
    for r in para.runs:
        r.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p('New Home Solutions', bold=True, size=32, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Technical Documentation', bold=True, size=18, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
p('Platform: Salesforce Lightning (LWC + Apex)', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_paragraph()
doc.add_paragraph()
p('Client: New Home Solutions', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Development Partner: CRM Mates Ltd, London', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Lead Salesforce Consultant: Deepak K Rana', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
p('deepak@crmmates.com | 07443 340401', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_paragraph()
p('Last Updated: 7 April 2026', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture',
    '3. NHS Process Pipeline',
    '4. Custom Objects & Fields',
    '5. LWC Component Catalogue',
    '6. Apex Class Catalogue',
    '7. External Integrations',
    '8. Final Checks Module',
    '9. Communications Hub',
    '10. Deployment & Environments',
    '11. Known Issues & Workarounds',
    '12. Change Log',
]
for item in toc_items:
    bullet(item)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'New Home Solutions (NHS) is a property sales and valuation management platform built entirely on '
    'Salesforce Lightning. The platform manages the end-to-end lifecycle of property applications — from '
    'initial submission through agent valuations, final checks, and vendor discussions to completion and commission.'
)
doc.add_paragraph(
    'CRM Mates Ltd is the development partner responsible for the design, build, and ongoing enhancement of this platform.'
)

doc.add_heading('Key Statistics', level=2)
add_table(
    ['Metric', 'Count'],
    [
        ['Lightning Web Components', '107+'],
        ['Apex Classes (production)', '77'],
        ['Apex Test Classes', '30+'],
        ['External Integrations', '5'],
        ['Opportunity Custom Fields', '192'],
        ['Record Types (Opportunity)', '13'],
        ['NHS Process Stages', '9'],
    ]
)

doc.add_heading('Business Context', level=2)
doc.add_paragraph(
    'New Home Solutions acts as an intermediary between house builders and property vendors. The platform enables NHS staff to:'
)
bullet('Receive and manage property applications from house builders')
bullet('Schedule vendor availability for agent valuations')
bullet('Book up to 3 estate agents per property for independent valuations')
bullet('Track valuation figures and generate NHS recommended pricing')
bullet('Perform final quality checks before vendor discussions')
bullet('Communicate with house builders via integrated email')
bullet('Generate and distribute PDF reports via Dropbox')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture', level=1)

doc.add_heading('Technology Stack', level=2)
add_table(
    ['Layer', 'Technology'],
    [
        ['Frontend', 'Lightning Web Components (LWC)'],
        ['Backend', 'Apex (Controllers, Services, Batch, Queueable)'],
        ['PDF Generation', 'Visualforce Pages rendered to PDF'],
        ['File Storage', 'Dropbox (OAuth2 integration)'],
        ['Address Lookup', 'Ideal Postcodes API'],
        ['Property Data', 'Street Data API, UK Property Data (RapidAPI)'],
        ['Maps', 'Google Maps Street View API'],
        ['Email', 'Salesforce EmailMessage + Org-Wide Email Address'],
        ['API', 'Apex @RestResource endpoint (OpportunityService)'],
    ]
)

doc.add_heading('Design Patterns', level=2)
bullet('Service Layer Architecture — Business logic in service classes (DropboxFileService, VendorAvailabilityService, PropertyReportService)')
bullet('Controller Pattern — Each LWC feature area has a dedicated @AuraEnabled Apex controller')
bullet('Batch Processing — OpportunityStageUpdaterBatch for automated stage transitions')
bullet('Queueable Jobs — TinyURLShortenerQueueable for async URL shortening')
bullet('Future Methods — @future(callout=true) in DropboxFileService for non-blocking API callouts')
bullet('Custom LWC Services — Reusable service components: toastService, scrollService, filterService, configService, dragDropService')

doc.add_heading('Data Model Overview', level=2)
doc.add_paragraph(
    'The platform is Opportunity-centric. Each property application is an Opportunity record with 192 custom fields, '
    'lookups to Account (House Builder, Agents 1-3) and Contact (Vendors 1-2), and related objects: '
    'Vendor_Availability__c, Vendor_Note__c, NHS_Property__c. API credentials are stored in DropBox__mdt custom metadata.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 3. NHS PROCESS PIPELINE
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. NHS Process Pipeline', level=1)
doc.add_paragraph(
    'The NHS_Process__c picklist field on Opportunity drives the entire application workflow. '
    'It is a restricted picklist assigned to all 13 record types.'
)

doc.add_heading('Stage Definitions', level=2)
add_table(
    ['#', 'Stage (API Value)', 'Label', 'Description'],
    [
        ['1', 'Application', 'Applications', 'New application received. Vendor and property details captured. Default stage.'],
        ['2', 'Vendor Availability', 'Vendor Availability', 'Vendor provides calendar availability slots for agent visits.'],
        ['3', 'Agents Booked', 'Book Agents', 'Up to 3 agents booked for property valuations.'],
        ['4', 'Figures to chased', 'Figures to Chase', 'Auto-triggered when appointment dates pass without valuations received.'],
        ['5', 'Valuations Ready', 'Valuations Ready', 'All agent valuation forms completed and PDF reports generated.'],
        ['6', 'Figures returned', 'Figures Returned', 'All 3 agents have submitted initial, target, and bottom line prices.'],
        ['7', 'Final Checks', 'Final Checks', 'Expert cross-check of all application data before vendor discussions.'],
        ['8', 'Vendor Discussions', 'Vendor Discussions', 'Direct discussions with vendors to finalise terms.'],
        ['9', 'Archived', 'Archived', 'Completed or cancelled applications.'],
    ]
)

doc.add_heading('Stage Gate Validations', level=2)
add_table(
    ['Transition', 'Validation', 'Error Message'],
    [
        ['Any → Agents Booked', 'At least 1 vendor availability slot from tomorrow onwards', 'At least 1 vendor availability slot is required before moving to Book Agents.'],
        ['Any → Final Checks', 'All 12 financial fields must be > 0: Agent 1/2/3 (Initial, Target, Bottom) + NHS (Market, Target, Forced)', 'All financial values must be entered before moving to Final Checks.'],
    ]
)

doc.add_heading('Automated Stage Transitions', level=2)
doc.add_paragraph(
    'OpportunityStageUpdaterBatch — Scheduled batch job queries Opportunities in "Agents Booked" stage. '
    'If appointment dates (Agent_1/2/3_Appointment__c) have passed and valuation not received, '
    'automatically updates NHS_Process__c to "Figures to Chased".'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 4. CUSTOM OBJECTS & FIELDS
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Custom Objects & Fields', level=1)

doc.add_heading('Opportunity — Application Details', level=2)
add_table(
    ['Field API Name', 'Type', 'Purpose'],
    [
        ['Application_Reference_Number__c', 'Text', 'Auto-generated reference'],
        ['Date_of_Application_Received__c', 'Date', 'When application was received'],
        ['Development__c', 'Text', 'Development/site name'],
        ['Plot__c', 'Text', 'Plot number'],
        ['Property_Address__c', 'Text', 'Full property address'],
        ['Scheme__c', 'Picklist', 'Assisted Sale / New Home / Part Exchange'],
        ['NHS_Process__c', 'Picklist (Restricted)', 'Current pipeline stage'],
    ]
)

doc.add_heading('Opportunity — Agent Assignments (per agent, x3)', level=2)
add_table(
    ['Field API Name', 'Type', 'Purpose'],
    [
        ['Agent_X__c', 'Account Lookup', 'Estate agent company'],
        ['Agent_X_Appointment__c', 'DateTime', 'Scheduled appointment'],
        ['Agent_X_Initial_Asking_Price__c', 'Currency', 'Agent initial asking price'],
        ['Agent_X_Target_Sale__c', 'Currency', 'Agent target sale price'],
        ['Agent_X_Bottom_Line__c', 'Currency', 'Agent bottom line price'],
        ['Agent_X_Valuation_Recieved__c', 'Checkbox', 'Whether valuation received'],
        ['Agent_X_Desktop_Valuation__c', 'Checkbox', 'Desktop-only valuation (no visit)'],
    ]
)

doc.add_heading('Opportunity — NHS Recommended Pricing', level=2)
add_table(
    ['Field API Name', 'Type', 'Purpose'],
    [
        ['Current_Asking_Price__c', 'Currency', 'NHS recommended market value'],
        ['Target_Sale__c', 'Currency', 'NHS recommended target sale'],
        ['Forced_Sale__c', 'Currency', 'NHS recommended forced sale'],
    ]
)

doc.add_heading('Opportunity — Final Checks (FC_ prefix)', level=2)
add_table(
    ['Field API Name', 'Type', 'Purpose'],
    [
        ['FC_Agent_1_Report__c', 'Checkbox', 'Agent 1 report verified'],
        ['FC_Agent_2_Report__c', 'Checkbox', 'Agent 2 report verified'],
        ['FC_Agent_3_Report__c', 'Checkbox', 'Agent 3 report verified'],
        ['FC_NHS_Pre_Will_Report__c', 'Checkbox', 'NHS Pre-Will report verified'],
        ['FC_Will_Report__c', 'Checkbox', 'Will report verified'],
        ['FC_Photos_Validated__c', 'Checkbox', 'Property photos validated'],
        ['FC_Address_Validated__c', 'Checkbox', 'Property address validated'],
    ]
)

doc.add_heading('Opportunity — Kanban Management', level=2)
add_table(
    ['Field API Name', 'Type', 'Purpose'],
    [
        ['Pinned__c', 'Checkbox', 'Card pinned to top of kanban column (max 3)'],
        ['SortOrder__c', 'Number', 'Custom sort order within kanban column'],
        ['Archived__c', 'Checkbox', 'Application archived'],
        ['Blacklisted__c', 'Checkbox', 'Blacklisted flag'],
    ]
)

doc.add_heading('Vendor_Availability__c', level=2)
doc.add_paragraph('Stores vendor availability for agent booking. Auto-number name: VA-{0000}.')
add_table(
    ['Field', 'Type', 'Purpose'],
    [
        ['Date__c', 'Date', 'Availability date'],
        ['Vendor__c', 'Contact Lookup', 'Vendor contact'],
        ['Agent__c', 'Account Lookup', 'Agent account'],
        ['Application__c', 'Opportunity Lookup', 'Related application'],
        ['AM__c / PM__c', 'Checkbox', 'AM/PM availability toggle'],
        ['Hour_00__c – Hour_23__c', 'Checkbox (x24)', 'Hourly slot availability'],
        ['Booked__c', 'Checkbox', 'Slot has been booked'],
    ]
)

doc.add_heading('Vendor_Note__c', level=2)
doc.add_paragraph('Communication notes linked to applications. Auto-number name: VN-{0000}.')
add_table(
    ['Field', 'Type', 'Purpose'],
    [
        ['Application__c', 'Opportunity Lookup', 'Related application'],
        ['Vendor__c', 'Contact Lookup', 'Related vendor'],
        ['Notes__c', 'Long Text', 'Note content'],
    ]
)

doc.add_heading('NHS_Property__c', level=2)
add_table(
    ['Field Group', 'Fields'],
    [
        ['Address', 'Address__Street__s, Address__City__s, Address__PostalCode__s'],
        ['Type', 'Detached, Semi_Detached, End_Terrace, Mid_Terrace, Apartment, Maisonette, Bungalow'],
        ['Features', 'Number_Of_Bedrooms, Front_Garden, Back_Garden, Parking, Garage, Extension'],
        ['Tenure', 'Freehold, Lease_Hold, Years_left_of_Lease, Service_Charge, Ground_Rent'],
        ['Other', 'Type_of_Heating, Building_Regs_Planning_Permission, Principle_Residence'],
    ]
)

doc.add_heading('Record Types (Opportunity)', level=2)
add_table(
    ['Record Type', 'Purpose'],
    [
        ['New_Home_Sales', 'Primary record type for new home sales applications'],
        ['Assisted_Sales', 'Assisted sale scheme applications'],
        ['Part_Exchange_Applications', 'Part exchange initial applications'],
        ['Part_Exchange_Valuations_In_Process', 'Part exchange valuations in progress'],
        ['Part_Exchange_Availability_Report', 'Part exchange availability reports'],
        ['Part_Exchange_Offers', 'Part exchange offers'],
        ['Part_Exchange_Instruction', 'Part exchange instructions'],
        ['Part_Exchange_Sales_Progression', 'Part exchange sales progression'],
        ['Part_Exchange_Commission', 'Part exchange commission tracking'],
        ['Part_Exchange_Completed', 'Completed part exchanges'],
        ['Events_and_Support', 'Events and support activities'],
        ['Archived', 'Archived/completed applications'],
        ['Unknown', 'Default/unclassified'],
    ]
)
doc.add_paragraph('All 13 record types have the full set of 9 NHS_Process__c picklist values assigned.')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 5. LWC COMPONENT CATALOGUE
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. LWC Component Catalogue', level=1)

doc.add_heading('Kanban & Pipeline Management', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['nhsApplicationKanbanV7', 'Main kanban board with 9-stage navigation, drag-drop, pin/archive, filtering', 'NhsApplicationKanbanController'],
        ['allApplicationKanban', 'All applications kanban view', 'AllApplicationKanban'],
        ['archivedApplicationKanban', 'Archived application management', 'ArchivedApplicationKanbanController'],
        ['mainKanbanScreen', 'Parent container for kanban displays', '—'],
    ]
)

doc.add_heading('Application Forms', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['nhsApplicationForm', 'Main NHS application form (Create Application page)', 'ApplicationFormController'],
        ['nhsApplicationFormPro', 'Professional version of application form', '—'],
        ['houseBuilderApplicationForm', 'House builder specific application form', 'houseBuilderApplication'],
    ]
)

doc.add_heading('View Application (Record Page)', level=2)
add_table(
    ['Component', 'Purpose'],
    [
        ['nhsOpportunityDetailedView', 'Full opportunity record page with stage tracker, property details, agent valuations, conditional section rendering'],
    ]
)

doc.add_heading('Vendor Availability & Scheduling', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['nhsVendorAvailability', 'Calendar interface with hourly slots, AM/PM toggles, week navigation', 'VendorAvailabilityService'],
        ['nhsVendorAvailabilityList', 'List of vendor availability records', 'VendorAvailabilityListController'],
        ['nhsAgentBooking', 'Agent booking interface', 'EventController'],
        ['nhsAgentBookedList', 'List of agents with booked appointments', 'AgentBookedListController'],
    ]
)

doc.add_heading('Final Checks', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['nhsFinalChecksList', 'Kanban stage view — table with inline checklist expansion', 'FinalChecksController'],
        ['nhsFinalChecksPanel', 'Lightweight inline panel for record pages', 'FinalChecksController'],
        ['nhsFinalChecksPage', 'Full record page with read-only details, agent valuations, checks, and email', 'FinalChecksController, NHSCommunicationsController'],
    ]
)

doc.add_heading('Communications & Documents', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['nhsCommunicationsHub', 'Central email and call management dashboard', 'NHSCommunicationsController'],
        ['nhsDropboxBrowser', 'Dropbox file browser with folder navigation, upload, download', 'DropboxBrowserController'],
        ['nhsDropboxSetup', 'Dropbox OAuth configuration', 'DropboxOAuthController'],
    ]
)

doc.add_heading('Chasing & Follow-ups', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['chasingSystem / chasingParentComponent', 'Follow-up workflow for overdue valuations', 'EventController'],
        ['nhsFiguresToChaseList', 'Outstanding figures requiring follow-up', 'FiguresToChaseController'],
        ['nhsFiguresReturnedList', 'Valuations returned tracking', 'FiguresReturnedController'],
        ['nhsValuationsReadyList', 'Valuations ready for review', 'ValuationsReadyController'],
    ]
)

doc.add_heading('Address & Property Lookup', level=2)
add_table(
    ['Component', 'Purpose', 'Apex Dependency'],
    [
        ['idealPostcodeAutocomplete', 'Ideal Postcodes API address autocomplete', 'IdealPostcodeController'],
        ['propertySearchAPI', 'Street Data API for EPC ratings', 'PropertySearchAPIController'],
        ['mapViewer', 'Property map visualisation', '—'],
        ['callStreetAPI', 'Google Street View images', 'StreetViewController'],
    ]
)

doc.add_heading('Utility Components', level=2)
add_table(
    ['Component', 'Purpose'],
    [
        ['customLookup / customLookupLwc', 'Reusable lookup/search fields'],
        ['customLoader', 'Loading spinner'],
        ['customCarousel / lightbox', 'Image carousel and lightbox viewer'],
        ['toastService / scrollService / filterService', 'Shared LWC service components'],
        ['configService / dragDropService', 'Configuration and drag-drop services'],
        ['pagination2', 'Pagination'],
        ['initialsCircle / accountCircles', 'Avatar/initials display'],
        ['dateRangeSlider', 'Date range filtering'],
        ['ePCRating / dynamicEPC', 'Energy Performance Certificate display'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. APEX CLASS CATALOGUE
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Apex Class Catalogue', level=1)

doc.add_heading('Application & Opportunity Controllers', level=2)
add_table(
    ['Class', 'Sharing', 'Purpose'],
    [
        ['NhsApplicationKanbanController', 'with sharing', 'Kanban data retrieval, filtering, sorting, pin/archive, card ordering'],
        ['AllApplicationKanban', 'with sharing', 'All applications kanban data'],
        ['ApplicationFormController', '—', 'Application form data management'],
        ['OpportunityController', '—', 'General opportunity operations'],
        ['OpportunityService', '—', '@RestResource endpoint for external opportunity creation'],
    ]
)

doc.add_heading('Final Checks', level=2)
add_table(
    ['Class', 'Sharing', 'Methods'],
    [
        ['FinalChecksController', 'with sharing', 'getFinalChecksApplications, getFinalChecks, getFinalChecksPageData, saveFinalChecks'],
    ]
)

doc.add_heading('Vendor & Agent Management', level=2)
add_table(
    ['Class', 'Purpose'],
    [
        ['VendorAvailabilityService', 'Complex availability management: hourly slots, AM/PM, event creation, booking validation'],
        ['VendorAvailabilityListController', 'List view of availability records'],
        ['EventController', 'Event/appointment management, chasing count, valuation tracking'],
        ['VendorNoteController', 'Vendor note CRUD'],
    ]
)

doc.add_heading('Valuation & Figures', level=2)
add_table(
    ['Class', 'Sharing', 'Purpose'],
    [
        ['FiguresToChaseController', 'with sharing', 'Outstanding valuations list'],
        ['FiguresReturnedController', 'with sharing', 'All-figures-returned applications list'],
        ['ValuationsReadyController', 'with sharing', 'Valuations ready for processing'],
    ]
)

doc.add_heading('Communications', level=2)
add_table(
    ['Class', 'Methods'],
    [
        ['NHSCommunicationsController', 'getCommunications, getEmailTemplates, getRenderedTemplate, sendEmail, sendEmailWithBcc, logCall'],
    ]
)

doc.add_heading('Dropbox Integration', level=2)
add_table(
    ['Class', 'Purpose'],
    [
        ['DropboxFileService', 'Core Dropbox API: upload, shared links, folder management. @future(callout=true)'],
        ['DropboxOAuthController', 'OAuth2 token exchange and refresh'],
        ['DropboxBrowserController', 'List folders/files, metadata display'],
    ]
)

doc.add_heading('PDF Generation', level=2)
add_table(
    ['Class', 'Purpose'],
    [
        ['PdfGeneratorController', 'Application PDF generation and attachment'],
        ['HouseBuilderPdfController', 'House builder application PDFs'],
        ['valuationFormPdfController', 'Valuation assessment report PDFs'],
        ['SendPDFonCreationHandler', 'Trigger handler: auto-sends valuation PDF via email'],
    ]
)

doc.add_heading('Batch & Scheduled Jobs', level=2)
add_table(
    ['Class', 'Purpose'],
    [
        ['OpportunityStageUpdaterBatch', 'Queries "Agents Booked", checks past appointments, auto-updates to "Figures to Chased"'],
        ['OpportunityStageUpdaterSchedule', 'Scheduler for the batch job'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 7. EXTERNAL INTEGRATIONS
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. External Integrations', level=1)

doc.add_heading('Dropbox', level=2)
add_table(
    ['Detail', 'Value'],
    [
        ['Auth', 'OAuth2 with refresh token flow'],
        ['Upload Endpoint', 'https://content.dropboxapi.com/2/files/upload'],
        ['Shared Links', 'https://api.dropboxapi.com/2/sharing/create_shared_link_with_settings'],
        ['Credentials', 'Stored in DropBox__mdt custom metadata'],
        ['Folder Structure', '/Private Salesforce Documents/Accounts/[AccountName]/Logo|Application|Valuation Assessment Report/'],
        ['Processing', '@future(callout=true) for non-blocking uploads'],
    ]
)

doc.add_heading('Ideal Postcodes', level=2)
add_table(
    ['Detail', 'Value'],
    [
        ['Endpoint', 'https://api.ideal-postcodes.co.uk/v1/addresses'],
        ['API Key', 'Stored in Custom Label (System.Label.idealpcapi)'],
        ['Purpose', 'Address autocomplete and postcode lookup'],
    ]
)

doc.add_heading('Street Data API', level=2)
add_table(
    ['Detail', 'Value'],
    [
        ['Endpoint', 'https://api.data.street.co.uk/street-data-api/v2/'],
        ['Purpose', 'EPC ratings and energy performance data by postcode'],
    ]
)

doc.add_heading('UK Property Data (RapidAPI)', level=2)
add_table(
    ['Detail', 'Value'],
    [
        ['Endpoint', 'https://api-ir7ctmwisa-ew.a.run.app/'],
        ['Method', 'propertytools.api.v1.Public/GetPropertyReport'],
        ['Purpose', 'Detailed property reports by postcode/PAON'],
    ]
)

doc.add_heading('Google Maps', level=2)
add_table(
    ['Detail', 'Value'],
    [
        ['Endpoint', 'https://maps.googleapis.com/maps/api/streetview/metadata'],
        ['Purpose', 'Street View imagery and property visualisation'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 8. FINAL CHECKS MODULE
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. Final Checks Module', level=1)
doc.add_paragraph(
    'The Final Checks module provides expert quality assurance before an application moves to Vendor Discussions. '
    'It consists of 7 validation checkpoints, read-only financial summaries, and integrated email to notify the house builder.'
)

doc.add_heading('Data Model', level=2)
doc.add_paragraph('Seven checkbox fields on Opportunity, all prefixed FC_:')
add_table(
    ['Field', 'Label', 'Check Group'],
    [
        ['FC_Agent_1_Report__c', 'Agent 1 Report', 'Agent Reports'],
        ['FC_Agent_2_Report__c', 'Agent 2 Report', 'Agent Reports'],
        ['FC_Agent_3_Report__c', 'Agent 3 Report', 'Agent Reports'],
        ['FC_NHS_Pre_Will_Report__c', 'NHS Pre-Will Report', 'Will Reports'],
        ['FC_Will_Report__c', 'Will Report', 'Will Reports'],
        ['FC_Photos_Validated__c', 'Photos Validated', 'Validation'],
        ['FC_Address_Validated__c', 'Address Validated', 'Validation'],
    ]
)

doc.add_heading('Apex Controller: FinalChecksController', level=2)
add_table(
    ['Method', 'Cacheable', 'Purpose'],
    [
        ['getFinalChecksApplications(searchTerm)', 'No', 'Queries all applications in "Final Checks" stage with search'],
        ['getFinalChecks(opportunityId)', 'No', 'Returns Map<String, Boolean> of 7 check values'],
        ['getFinalChecksPageData(opportunityId)', 'No', 'Full page data: app details, 3 agents valuations, NHS pricing, checks, builder contacts'],
        ['saveFinalChecks(opportunityId, checks)', 'No', 'Updates 7 checkbox fields from Map<String, Boolean>'],
    ]
)

doc.add_heading('LWC Components', level=2)

p('nhsFinalChecksList (Kanban Stage View)', bold=True, size=11)
bullet('Displayed when user clicks "Final Checks" tab in the kanban')
bullet('Table of all applications with progress bar (X/7) and status badge')
bullet('Click row to expand inline checklist panel with grouped checkboxes')
bullet('Save/Cancel per application, search and refresh')

p('nhsFinalChecksPage (Record Page View)', bold=True, size=11)
doc.add_paragraph('Displayed on the Opportunity detail page when NHS_Process__c = "Final Checks". Hides all other sections.')
doc.add_paragraph('Sections displayed:')
bullet('Application Details (read-only) — Application Name, Builder, Vendor, Email, Mobile')
bullet('Agent Valuations (read-only) — 3-agent table with Initial/Target/Bottom + Average. Agent headers colour-coded: Agent 1 green, Agent 2 blue, Agent 3 purple')
bullet('NHS Recommendation (read-only) — Market Value, Target Sale, Forced Sale')
bullet('Final Checks (editable) — 7 checkboxes in 3 groups with progress bar')
bullet('Final Email (conditional) — Unlocked only when all 7 checks complete')

doc.add_heading('Email Workflow', level=2)
bullet('Template picker opens in a lightbox modal with search')
bullet('Selected template rendered with merge fields via NHSCommunicationsController')
bullet('Recipients are builder contacts (auto-populated from House_Builder__c Account Contacts with email)')
bullet('All contacts pre-selected with toggle pills; manual CC and BCC fields available')
bullet('Subject line editable; email body preview rendered as HTML')
bullet('"Confirm & Send" button — first selected contact as To, remaining as CC, manual BCC as BCC')

doc.add_heading('Stage Gate Validation', level=2)
doc.add_paragraph(
    'Transition to "Final Checks" requires all 12 financial fields to have values > 0: '
    'Agent 1/2/3 (Initial, Target, Bottom) + NHS (Market, Target, Forced). '
    'Validation is enforced in nhsOpportunityDetailedView.handleStageChange() with a sticky error toast on failure.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 9. COMMUNICATIONS HUB
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Communications Hub', level=1)
doc.add_paragraph(
    'The Communications Hub (nhsCommunicationsHub LWC + NHSCommunicationsController Apex) provides '
    'centralised email and call management per Opportunity.'
)

doc.add_heading('Capabilities', level=2)
add_table(
    ['Feature', 'Detail'],
    [
        ['Email History', 'Queries EmailMessage records — sender, recipient, subject, status, timestamps'],
        ['Call Logging', 'Creates Task records with CallType, duration, status'],
        ['Email Templates', 'Loads from "NHS Email Templates" folder. Dynamic merge field rendering'],
        ['Outbound Email', 'Sends via Org-Wide Email Address: systemtest@newhomesolutions.co.uk'],
        ['CC/BCC Support', 'sendEmailWithBcc method parses comma-separated CC and BCC lists'],
        ['Attachments', 'PDF attachments from Dropbox-hosted files'],
    ]
)

doc.add_heading('Apex Methods', level=2)
add_table(
    ['Method', 'Purpose'],
    [
        ['getCommunications(opportunityId)', 'Returns email + call history'],
        ['getEmailTemplates()', 'Lists active templates from NHS folder'],
        ['getRenderedTemplate(templateId, opportunityId)', 'Renders template with merge fields'],
        ['sendEmail(...)', 'Basic email send (delegates to sendEmailWithBcc)'],
        ['sendEmailWithBcc(...)', 'Full email send with CC/BCC parsing'],
        ['logCall(...)', 'Log call activity as Task'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT & ENVIRONMENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. Deployment & Environments', level=1)

doc.add_heading('Target Orgs', level=2)
add_table(
    ['Org', 'Alias', 'Username'],
    [
        ['Training Sandbox', 'SandboxDev', 'deepak-nhs-ee@crmmates.com.training'],
    ]
)

doc.add_heading('Common CLI Commands', level=2)
p('Deploy specific components:', bold=True, size=10)
doc.add_paragraph(
    'sf project deploy start --source-dir force-app/main/default/lwc/componentName '
    '--source-dir force-app/main/default/classes/ClassName.cls --target-org SandboxDev'
).style = 'No Spacing'

p('Deploy all record types:', bold=True, size=10)
doc.add_paragraph(
    'sf project deploy start --source-dir force-app/main/default/objects/Opportunity/recordTypes --target-org SandboxDev'
).style = 'No Spacing'

p('Check deploy status (when CLI shows locale error):', bold=True, size=10)
doc.add_paragraph(
    'sf project deploy report --job-id <DEPLOY_ID> --target-org SandboxDev'
).style = 'No Spacing'

doc.add_heading('Project Structure', level=2)
structure = [
    'force-app/main/default/',
    '  classes/               — Apex classes + meta XML',
    '  lwc/                   — Lightning Web Components (107+)',
    '  objects/Opportunity/',
    '    fields/              — 192 custom field definitions',
    '    recordTypes/         — 13 record types',
    '    validationRules/     — 4 validation rules (all inactive)',
    '  pages/                 — Visualforce pages (PDF)',
    '  triggers/              — Apex triggers',
]
for line in structure:
    doc.add_paragraph(line).style = 'No Spacing'
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 11. KNOWN ISSUES & WORKAROUNDS
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Known Issues & Workarounds', level=1)
add_table(
    ['Issue', 'Detail', 'Workaround'],
    [
        ['CLI Locale Error', 'Missing message metadata.transfer:Finalizing for locale en_US on deploy', 'Deploy succeeds despite error. Use sf project deploy report --job-id <ID> to verify'],
        ['Picklist Typo', 'Figures to chased (API value) vs Figures to chase (label)', 'Maintained for backward compatibility — all code references the API value'],
        ['FLS on New Fields', 'Newly deployed custom fields may not be visible via Lightning Data Service', 'Use Apex controllers instead of LDS for new fields'],
        ['Restricted Picklist + Record Types', 'Adding new picklist values requires updating ALL 13 record types', 'Always deploy record type changes alongside picklist changes'],
    ]
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 12. CHANGE LOG
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. Change Log', level=1)
add_table(
    ['Date', 'Change', 'Components Affected'],
    [
        ['2026-04-07', 'Added Final Checks and Vendor Discussions stages (7 & 8)', 'NHS_Process__c, 13 record types, nhsApplicationKanbanV7, nhsOpportunityDetailedView'],
        ['2026-04-07', 'Created 7 Final Checks checkbox fields', 'FC_Agent_1/2/3_Report, FC_NHS_Pre_Will_Report, FC_Will_Report, FC_Photos/Address_Validated'],
        ['2026-04-07', 'Built FinalChecksController Apex class', 'getFinalChecksApplications, getFinalChecks, getFinalChecksPageData, saveFinalChecks'],
        ['2026-04-07', 'Built 3 Final Checks LWC components', 'nhsFinalChecksList, nhsFinalChecksPanel, nhsFinalChecksPage'],
        ['2026-04-07', 'Added stage gate validation for Final Checks', 'nhsOpportunityDetailedView — validates 12 financial fields'],
        ['2026-04-07', 'Added sendEmailWithBcc to NHSCommunicationsController', 'Supports comma-separated CC and BCC lists'],
        ['2026-04-07', 'Updated View Application: hide sections during Final Checks', 'Property Details, Agent Details, Dropbox, Vendor Notes hidden'],
        ['2026-04-07', 'Improved fonts and currency inputs', 'Consistent 12.5px fonts, no spinners, centered £ values'],
        ['2026-04-07', 'Created project presentation', 'presentation.html — 15 slides, NHS brand, 9-stage pipeline'],
    ]
)

# ── Footer ──
doc.add_paragraph()
doc.add_paragraph()
p('Document maintained by: Deepak K Rana, Lead Salesforce Consultant, CRM Mates Ltd', size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
p('deepak@crmmates.com | 07443 340401', size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
output = os.path.join(os.path.dirname(__file__), 'TECHNICAL_DOCUMENT.docx')
doc.save(output)
print(f'Generated: {output}')
